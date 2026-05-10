"""Build BEFORE_AFTER_COMPARISON.docx — landscape 3-col table showing
each section's original vs revised body text side by side."""
import os, sys, re, difflib
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), '..',
                    'docs', 'publication', 'submissions')
ROOT = os.path.abspath(ROOT)

ORIG = os.path.join(ROOT, 'jlp_initial_2026-03-31',  'MANUSCRIPT (1).docx')
REV  = os.path.join(ROOT, 'jlp_revision_2026-05-08', 'PAPER_JLP_REVISED_v3.docx')
OUT  = os.path.join(ROOT, 'jlp_revision_2026-05-08', 'BEFORE_AFTER_COMPARISON.docx')

# Truncate per-section excerpt to keep the table readable
EXCERPT_CHARS = 1200


def parse_sections(path):
    """Return list of dicts: {level, num, title, heading_full, body}."""
    d = Document(path)
    sections = []
    cur = None
    for p in d.paragraphs:
        style = (p.style.name or '').lower() if p.style else ''
        if style.startswith('heading') or style == 'title':
            if cur is not None and (cur['body'] or cur['title']):
                sections.append(cur)
            level = 0
            m = re.match(r'heading\s*(\d+)', style)
            if m: level = int(m.group(1))
            txt = p.text.strip()
            num_m = re.match(r'^([\d\.A-Z]+(?:\.\d+)*)\.?\s+(.*)$', txt)
            if num_m:
                num = num_m.group(1).rstrip('.')
                title = num_m.group(2).strip()
            else:
                num, title = '', txt
            cur = {'level': level, 'num': num, 'title': title,
                   'heading_full': txt, 'body': ''}
        else:
            text = p.text.strip()
            if text and cur is not None:
                cur['body'] += (' ' if cur['body'] else '') + text
    if cur is not None:
        sections.append(cur)
    return sections


def normalise(s):
    s = re.sub(r'[^\w\s]', ' ', s.lower())
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def best_section_match(rev_sec, orig_secs, used):
    rev_title = normalise(rev_sec['title'])
    if not rev_title:
        return -1, 0.0
    best, best_score = -1, 0.0
    for i, o in enumerate(orig_secs):
        if i in used:
            continue
        score = difflib.SequenceMatcher(None, rev_title,
                                        normalise(o['title'])).ratio()
        # Bonus for same section number
        if rev_sec['num'] and o['num']:
            if rev_sec['num'] == o['num']:
                score += 0.10
            elif rev_sec['num'].split('.')[0] == o['num'].split('.')[0]:
                score += 0.04
        if score > best_score:
            best, best_score = i, score
    return (best, best_score) if best_score >= 0.45 else (-1, best_score)


def truncate(s, n=EXCERPT_CHARS):
    if len(s) <= n:
        return s
    cut = s[:n].rsplit(' ', 1)[0]
    return cut + ' …[truncated]'


def shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tc_pr.append(shd)


def set_col_widths(table, widths_inches):
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); table._element.insert(0, tblPr)
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout'); tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    for tg in table._element.findall(qn('w:tblGrid')):
        table._element.remove(tg)
    grid = OxmlElement('w:tblGrid')
    for w in widths_inches:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 1440)))
        grid.append(gc)
    table._element.insert(list(table._element).index(tblPr) + 1, grid)
    for row in table.rows:
        for cell, w in zip(row.cells, widths_inches):
            cell.width = Inches(w)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(int(w * 1440)))
            tcW.set(qn('w:type'), 'dxa')


def main():
    print('Parsing original...')
    orig_secs = parse_sections(ORIG)
    print(f'  {len(orig_secs)} sections')
    print('Parsing revised...')
    rev_secs  = parse_sections(REV)
    print(f'  {len(rev_secs)} sections')

    # Skip the very-top title/author paragraph that's not really a section
    rev_secs = [s for s in rev_secs if s['title']]
    orig_secs = [s for s in orig_secs if s['title']]

    # Match
    used = set()
    rows = []  # (label, type, before, after)
    n_unchanged_dropped = 0
    for rs in rev_secs:
        idx, score = best_section_match(rs, orig_secs, used)
        label = (f'{rs["num"]} {rs["title"]}'.strip()) if rs['num'] else rs['title']
        if idx >= 0:
            used.add(idx)
            os_ = orig_secs[idx]
            old_label = (f'{os_["num"]} {os_["title"]}'.strip()) if os_['num'] else os_['title']
            sim = difflib.SequenceMatcher(None, os_['body'], rs['body']).ratio()
            if sim >= 0.95:
                n_unchanged_dropped += 1
                continue  # drop UNCHANGED rows per user request
            before = (f'(was: {old_label})\n\n' + os_['body']) if os_['title'] != rs['title'] else os_['body']
            after  = rs['body']
            rows.append((label, 'REVISED', before, after))
        else:
            rows.append((label, 'NEW', '— (new section, no equivalent in original)', rs['body']))
    # Original sections that were never matched -> deleted
    for i, os_ in enumerate(orig_secs):
        if i in used: continue
        if not os_['body'] and not os_['title']: continue
        old_label = (f'{os_["num"]} {os_["title"]}'.strip()) if os_['num'] else os_['title']
        rows.append((old_label, 'DELETED', os_['body'],
                     '— (removed in revision; content typically folded into another section)'))
    print(f'  Dropped {n_unchanged_dropped} UNCHANGED rows (>=0.95 similarity, no useful before/after to show)')

    # Build docx
    doc = Document()
    section = doc.sections[0]
    new_w, new_h = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = new_w
    section.page_height = new_h
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.2)
    section.right_margin  = Cm(1.2)

    title = doc.add_paragraph()
    run = title.add_run('Before / After section comparison — JLP-D-26-00414 (revision 2026-05-10)')
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    n_new = sum(1 for r in rows if r[1] == 'NEW')
    n_rev = sum(1 for r in rows if r[1] == 'REVISED')
    n_del = sum(1 for r in rows if r[1] == 'DELETED')
    sub.add_run(
        f'{len(rows)} changed rows: {n_new} NEW, {n_rev} REVISED, {n_del} DELETED. '
        f'(Sections >=95% identical between versions are not shown.) '
        f'Excerpts truncated to {EXCERPT_CHARS} chars per cell. '
        f'See PAPER_JLP_REVISED_v3.docx for the full revised manuscript.'
    ).italic = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '#'
    hdr[1].text = 'Section'
    hdr[2].text = 'Before (initial submission, 2026-03-31)'
    hdr[3].text = 'After (this revision)'
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs: r.bold = True
        shade_cell(c, 'D9E1F2')

    BG = {'NEW': 'E2F0D9', 'REVISED': 'FFF2CC',
          'UNCHANGED': 'FFFFFF', 'DELETED': 'FBE5D6'}

    for i, (label, chg, before, after) in enumerate(rows, 1):
        row = table.add_row()
        c0, c1, c2, c3 = row.cells
        c0.text = str(i)
        type_para = c1.add_paragraph()
        type_para.add_run(f'[{chg}] ').bold = True
        type_para.add_run(label)
        # remove default empty paragraph
        c1._tc.remove(c1.paragraphs[0]._p)
        c2.text = truncate(before)
        c3.text = truncate(after)
        for c in row.cells:
            shade_cell(c, BG.get(chg, 'FFFFFF'))
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    set_col_widths(table, [0.4, 2.0, 4.0, 4.0])

    if os.path.exists(OUT): os.remove(OUT)
    doc.save(OUT)
    print(f'\nSaved: {OUT} ({os.path.getsize(OUT):,} B)')
    print(f'  Rows: {len(rows)} (NEW={n_new}, REVISED={n_rev}, DELETED={n_del})')


if __name__ == '__main__':
    main()
