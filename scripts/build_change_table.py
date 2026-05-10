"""Build CHANGE_SUMMARY_TABLE.md / .docx for the JLP revision package.

Outputs include a quantitative scope header (computed from the original
and revised .docx files) and a fixed-width-column section change table.
The pandoc output is post-processed via python-docx so the table columns
are not equal-width (the default is unreadable when row 1 is "1" and
row 5 is a 30-word description).
"""
import os, sys, re, subprocess
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), '..',
                    'docs', 'publication', 'submissions')
ROOT = os.path.abspath(ROOT)

ORIG_DOCX = os.path.join(ROOT, 'jlp_initial_2026-03-31',  'MANUSCRIPT (1).docx')
REV_DOCX  = os.path.join(ROOT, 'jlp_revision_2026-05-08', 'PAPER_JLP_REVISED_v3.docx')
ARCHIVE   = os.path.join(ROOT, 'jlp_revision_2026-05-08', '_archive')
OUT_MD    = os.path.join(ARCHIVE, 'CHANGE_SUMMARY_TABLE.md')
OUT_DOCX  = os.path.join(ROOT, 'jlp_revision_2026-05-08', 'CHANGE_SUMMARY_TABLE.docx')

CHANGES = [
    ("Abstract", "REVISED", "R1.4, R2.1",
     "Reframed synthetic results as implementation verification; tightened retrieval claims to stat-significant K (Recall@1, MRR@10 with paired-bootstrap 95% CI)."),
    ("1. Introduction", "REVISED", "R1.1, R2.1",
     "Tightened motivation; introduced jurisdiction compliance gap terminology; clarified two-RQ scope."),
    ("2.1 AI Applications in Process Safety", "REVISED", "R2.x",
     "Updated literature framing for process-safety AI."),
    ("2.2 RAG in Technical and Regulatory Domains", "REVISED", "R2.x",
     "Revised RAG-in-regulatory-domains positioning."),
    ("2.3 Multi-Discipline Maintenance and AI", "REVISED", "R2.x",
     "Clarified multi-discipline asset-integrity context."),
    ("2.4 Verification Methodology — N-Version / K-Voting", "REVISED", "R3.1",
     "Strengthened N-version programming background relevant to K-voting rationale."),
    ("2.5 Korean PSM Regulation and KOSHA", "REVISED", "R2.3",
     "Sharpened Korean-jurisdiction regulatory background."),
    ("2.6 Position relative to HAZOP, RBI, Digital Twin", "NEW", "R2.2",
     "New subsection positioning the framework against HAZOP, RBI, and Digital Twin scope."),
    ("3. System Architecture", "REVISED", "R1.1",
     "Named deterministic state machine and orchestrator (pipeline.py); added Figure 1."),
    ("4.1 Corpus Construction", "REVISED", "R2.4",
     "Corrected corpus counts to 1,327 guides + 3,102 statutory rows -> 16,174 indexed entries (12 excluded for empty body)."),
    ("4.2 Retrieval and Generation Pipeline", "REVISED", "R1.1, R2.x",
     "Formalised concept-aware query builder, synonym expansion, OR-fallback; added Figure 2."),
    ("4.3 KOSHA Regulatory Grounding — Mandatory vs Guidance", "NEW", "R2.3",
     "New subsection separating mandatory law-articles from KOSHA guides for citation policy."),
    ("4.4 Why standard EPC workflows do not cover these obligations", "NEW", "R2.3",
     "New subsection explaining the gap between EPC-standards review and KOSHA PSM coverage."),
    ("5.1 Domain Calculation Engines", "REVISED", "R1.1",
     "Added Table 1 mapping seven engines -> standards -> source files."),
    ("5.2 Four-Layer Hybrid Verification Model", "REVISED", "R1.1, R3.1",
     "Documented Layer-4 thresholds (2% warning, 5% escalation), aligned with reverse_check.py."),
    ("5.3 Cross-Discipline Consistency Validator", "REVISED", "R1.1",
     "Documented ten-pair coupling set; Figure 3 added."),
    ("5.4 Why three paths and a 1% tolerance? — K-voting design rationale", "NEW", "R3.1",
     "New subsection giving design rationale for K-voting (3 paths, 1% tolerance)."),
    ("6.1 Calculation Accuracy on the Golden Dataset (implementation verification)", "REVISED", "R1.5",
     "Reframed 220-case headline as implementation verification, not predictive validation; added benchmark-construction subsection."),
    ("6.2 Seven-Discipline Pipeline Evaluation", "REVISED", "R1.x",
     "Clarified 43-scenario pipeline-eval scope vs 60-scenario ablation in §6.3."),
    ("6.3 Cross-Discipline Validator Ablation (RQ1)", "REVISED", "R1.1",
     "Updated headline to 26/60 (+0.4333, drift-corrected); added per-failure-mode 22/3/1 partition; Figure 4 added."),
    ("6.4 KOSHA RAG Regulatory Grounding — Three Cases (RQ2)", "REVISED", "R2.3",
     "Each case now identifies the specific mandatory KOSHA article surfaced (Article 256 / 266 / B-M-18)."),
    ("6.5 Curated Retrieval Benchmark (RQ2)", "REVISED", "R1.4",
     "Added paired-bootstrap 95% CI; honest disclosure that Recall@3/@5 CI includes zero; Figure 5 added."),
    ("6.6 Industry-Baseline Comparison and Internal Ablation", "NEW", "R1.1",
     "New subsection: industry baseline (ASME/API pass = compliance complete) 0/3 vs 3/3, plus four-layer ablation Table 7b."),
    ("6.6.3 Real-plant pipeline-execution evidence", "NEW", "R1.5",
     "Body-level summary of VES-REAL-001 anonymised cryogenic flare-drum case; Article 266 surfaced naturally."),
    ("6.7 Citation Traceability and Negative-Case Evidence", "NEW", "R1.5",
     "New subsection on citation-traceability precision (100% by construction) and PIP-GOLD-003 negative case."),
    ("7.1 The jurisdiction compliance gap", "REVISED", "R1.x, R2.1",
     "Sharpened gap definition tying calculation-only review to KOSHA PSM obligations."),
    ("7.2 Compliance co-pilot, not replacement", "NEW", "R2.2",
     "New subsection positioning the framework as a compliance co-pilot for HAZOP / RBI / Digital Twin / Regulatory RAG."),
    ("7.3 HAZOP scope clarification", "NEW", "R2.2",
     "New subsection: knowledge-based / AI-assisted, not a replacement for dynamic HAZOP."),
    ("7.4 Fitness-for-Service vs EPC standards", "NEW", "R2.5",
     "New subsection separating FFS workflows from EPC-standards verification scope."),
    ("7.5 Process-safety contribution", "REVISED", "R2.x",
     "Clarified the process-safety contribution narrative."),
    ("7.6 Reproducibility and explainability", "NEW", "R1.x, R2.x",
     "New subsection documenting end-to-end reproducibility (scripts/reproduce_all.py 13/13)."),
    ("8. Limitations and Future Work", "REVISED", "R1.5, R2.1",
     "Expanded from 3 -> 8 limitations including framework-vs-validated scope and benchmark-construction independence."),
    ("9. Conclusion", "REVISED", "R1.x",
     "Synthesised the revised contribution: integrated platform, four-layer verification, KOSHA RAG."),
    ("Acknowledgements", "REVISED", "blind-review",
     "Acknowledgements reworded for strict blind-review compliance (no direct reviewer tokens in body)."),
    ("Code and Data Availability", "NEW", "R1.x",
     "New section with AGPL-3.0 code/data release statement."),
    ("Acronyms", "NEW", "JLP-style",
     "New acronym glossary (JLP house style)."),
    ("Key Definitions", "NEW", "JLP-style",
     "New key-term definitions for clarity."),
    ("Appendix A. Real-Plant Data-Sheet Validation (VES-REAL-001)", "NEW", "R1.5",
     "New appendix: anonymised real-plant cryogenic flare-drum case with UG-27 calculation + KOSHA RAG retrieval (both probe variants)."),
    ("Appendix B — Extended validation evidence", "NEW", "R1.x, R2.x",
     "New appendix: per-discipline accuracy, Recall@K (K = 1..10), pipeline latency, retrieval failure inventory, full reproducibility manifest."),
]


def docx_stats(path):
    d = Document(path)
    body_paras = [p.text for p in d.paragraphs if p.text.strip()]
    table_paras = []
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        table_paras.append(p.text)
    all_text = ' '.join(body_paras + table_paras)
    import zipfile
    z = zipfile.ZipFile(path)
    xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
    return {
        'words':    len(re.findall(r'\b[\w-]+\b', all_text)),
        'paras':    len(body_paras),
        'tables':   len(d.tables),
        'figures':  len(re.findall(r'<pic:pic[\s>]', xml)),
        'headings': sum(1 for p in d.paragraphs
                        if p.style and p.style.name and
                        p.style.name.lower().startswith('heading')),
    }


def fmt_delta(before, after):
    delta = after - before
    pct = (delta / before * 100) if before else float('inf')
    sign = '+' if delta >= 0 else ''
    if before == 0:
        return f"{after:,} (new)"
    return f"{before:,} -> {after:,}  ({sign}{delta:,}, {sign}{pct:.0f}%)"


def main():
    s1 = docx_stats(ORIG_DOCX)
    s2 = docx_stats(REV_DOCX)
    n_new = sum(1 for c in CHANGES if c[1] == 'NEW')
    n_rev = sum(1 for c in CHANGES if c[1] == 'REVISED')
    all_trig = ' '.join(c[2] for c in CHANGES)
    r1 = len(re.findall(r'\bR1\b', all_trig))
    r2 = len(re.findall(r'\bR2\b', all_trig))
    r3 = len(re.findall(r'\bR3\b', all_trig))

    md = [
        '# Change Summary — JLP-D-26-00414 (revision 2026-05-10)',
        '',
        '## Quantitative scope of revision',
        '',
        '| Metric | Initial submission (2026-03-31) | Revised submission (this) | Change |',
        '|---|---:|---:|---|',
        f'| Word count | {s1["words"]:,} | {s2["words"]:,} | {fmt_delta(s1["words"], s2["words"])} |',
        f'| Sections (headings) | {s1["headings"]:,} | {s2["headings"]:,} | {fmt_delta(s1["headings"], s2["headings"])} |',
        f'| Body paragraphs | {s1["paras"]:,} | {s2["paras"]:,} | {fmt_delta(s1["paras"], s2["paras"])} |',
        f'| Tables | {s1["tables"]:,} | {s2["tables"]:,} | {fmt_delta(s1["tables"], s2["tables"])} |',
        f'| Figures | {s1["figures"]:,} | {s2["figures"]:,} | {fmt_delta(s1["figures"], s2["figures"])} |',
        '',
        f'**TL;DR.** The revised manuscript is **{s2["words"]/s1["words"]:.1f}x** the original word count, '
        f'adds **{s2["tables"]-s1["tables"]} new tables**, **{s2["figures"]-s1["figures"]} new figures**, '
        f'and **{s2["headings"]-s1["headings"]} new sections**. This is a major rewrite, not a typo pass.',
        '',
        '## Per-section change list',
        '',
        f'**{len(CHANGES)} sections changed: {n_new} NEW + {n_rev} REVISED.** '
        f'Reviewer coverage in this list: R1 in ~{r1} rows, R2 in ~{r2} rows, R3 in ~{r3} rows.',
        '',
        '| # | Section | Type | Triggered by | What changed |',
        '|---:|---|:--:|---|---|',
    ]
    for i, (sec, typ, trig, desc) in enumerate(CHANGES, 1):
        md.append(f'| {i} | {sec} | **{typ}** | {trig} | {desc} |')

    md.extend([
        '',
        '*See `RESPONSE_TO_REVIEWERS_v2.docx` for the full per-comment author response, '
        'and `PAPER_JLP_REVISED_v3_MARKED.docx` for the word-level visual diff.*',
    ])

    md_text = '\n'.join(md) + '\n'
    with open(OUT_MD, 'w', encoding='utf-8') as f:
        f.write(md_text)
    print(f'Wrote markdown: {OUT_MD} ({len(md_text):,} chars)')

    # Render via pandoc
    r = subprocess.run(['pandoc', OUT_MD, '-o', OUT_DOCX],
                       capture_output=True, text=True)
    if r.returncode != 0:
        print('pandoc error:', r.stderr[:500])
        sys.exit(1)
    print(f'Pandoc render: {OUT_DOCX} ({os.path.getsize(OUT_DOCX):,} B)')

    # Post-process: set explicit column widths so the table is readable
    # Total page width landscape ~10in or portrait ~6.5in. Use portrait widths.
    PORTRAIT_TABLE_INCHES = 6.4

    # Two tables in the doc:
    #   Table 0: 4-column quantitative-scope table
    #   Table 1: 5-column change-list table
    SCOPE_RATIOS  = [0.18, 0.22, 0.22, 0.38]                # 4 cols
    CHANGE_RATIOS = [0.04, 0.30, 0.07, 0.12, 0.47]          # 5 cols (#, Section, Type, Trig, What)

    def set_widths(table, ratios):
        # Set tblLayout to "fixed" so word respects our widths
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._element.insert(0, tblPr)
        layout = tblPr.find(qn('w:tblLayout'))
        if layout is None:
            layout = OxmlElement('w:tblLayout')
            tblPr.append(layout)
        layout.set(qn('w:type'), 'fixed')
        # Build/replace tblGrid
        for tg in table._element.findall(qn('w:tblGrid')):
            table._element.remove(tg)
        tblGrid = OxmlElement('w:tblGrid')
        for ratio in ratios:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(int(ratio * PORTRAIT_TABLE_INCHES * 1440)))
            tblGrid.append(gc)
        table._element.insert(list(table._element).index(tblPr) + 1, tblGrid)
        # Set per-cell widths
        for row in table.rows:
            for cell, ratio in zip(row.cells, ratios):
                cell.width = Inches(ratio * PORTRAIT_TABLE_INCHES)
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(int(ratio * PORTRAIT_TABLE_INCHES * 1440)))
                tcW.set(qn('w:type'), 'dxa')

    doc = Document(OUT_DOCX)
    if len(doc.tables) >= 2:
        set_widths(doc.tables[0], SCOPE_RATIOS)
        set_widths(doc.tables[1], CHANGE_RATIOS)
        print(f'Fixed column widths on table 0 ({len(SCOPE_RATIOS)} cols) '
              f'and table 1 ({len(CHANGE_RATIOS)} cols).')
    else:
        print(f'WARN: expected 2 tables, got {len(doc.tables)} - skipping width fix')
    doc.save(OUT_DOCX)
    print(f'Final: {OUT_DOCX} ({os.path.getsize(OUT_DOCX):,} B)')


if __name__ == '__main__':
    main()
