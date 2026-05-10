"""Build PAPER_JLP_REVISED_v3_MARKED.docx with NATIVE Word track-changes.

This produces real <w:ins> and <w:del> elements that Word recognises as
tracked changes (with the side panel, accept/reject buttons, author/date
attribution). NOT just visual styling — the actual OOXML revision marks.

Approach:
  1. Read original .docx paragraphs (text, including table cells).
  2. Walk revised .docx paragraphs.
  3. For each revised paragraph that does have a matching original paragraph
     (SequenceMatcher ratio >= 0.40), replace its runs with a native diff
     using <w:ins> for inserted words and <w:del>/<w:delText> for deleted
     words. Equal segments stay as plain runs.
  4. For revised paragraphs with no original match, wrap the entire text
     content in a single <w:ins> (Word treats this as a wholly inserted
     paragraph).
  5. Append a "DELETED FROM ORIGINAL" page-break section: each unmatched
     original paragraph becomes a new paragraph whose entire content is
     wrapped in <w:del>.

Notes:
  - Heading paragraphs are skipped from diff to keep the structure readable.
  - Tables are skipped from diff for the same reason (table content was
    largely added in revision; word-level diff inside tables is unreadable).
  - All <w:ins>/<w:del> elements use a single author "Yuyong Kim" and a
    single fixed date so accept/reject groups them sensibly in Word.
"""
import os, sys, difflib, time
from datetime import datetime
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), '..',
                    'docs', 'publication', 'submissions')
ROOT = os.path.abspath(ROOT)

ORIG_DOCX = os.path.join(ROOT, 'jlp_initial_2026-03-31',  'MANUSCRIPT (1).docx')
REV_DOCX  = os.path.join(ROOT, 'jlp_revision_2026-05-08', 'PAPER_JLP_REVISED_v3.docx')
OUT_DOCX  = os.path.join(ROOT, 'jlp_revision_2026-05-08', 'PAPER_JLP_REVISED_v3_MARKED.docx')

AUTHOR = 'Yuyong Kim'
DATE_ISO = '2026-05-10T00:00:00Z'

# Counter for w:id (must be unique within document)
_next_id = 1
def next_id():
    global _next_id
    _next_id += 1
    return _next_id


def make_run(text, *, run_props_template=None):
    """Create <w:r><w:t>text</w:t></w:r>; optionally clone rPr from a template run."""
    r = OxmlElement('w:r')
    if run_props_template is not None:
        rPr_src = run_props_template.find(qn('w:rPr'))
        if rPr_src is not None:
            from copy import deepcopy
            r.append(deepcopy(rPr_src))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def make_del_run(text, *, run_props_template=None):
    """Create <w:r><w:delText>text</w:delText></w:r> for use inside <w:del>."""
    r = OxmlElement('w:r')
    if run_props_template is not None:
        rPr_src = run_props_template.find(qn('w:rPr'))
        if rPr_src is not None:
            from copy import deepcopy
            r.append(deepcopy(rPr_src))
    dt = OxmlElement('w:delText')
    dt.set(qn('xml:space'), 'preserve')
    dt.text = text
    r.append(dt)
    return r


def make_ins(child_run):
    """Wrap a run inside <w:ins> with author/date attributes."""
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(next_id()))
    ins.set(qn('w:author'), AUTHOR)
    ins.set(qn('w:date'), DATE_ISO)
    ins.append(child_run)
    return ins


def make_del(child_run):
    """Wrap a run inside <w:del> with author/date attributes."""
    d = OxmlElement('w:del')
    d.set(qn('w:id'), str(next_id()))
    d.set(qn('w:author'), AUTHOR)
    d.set(qn('w:date'), DATE_ISO)
    d.append(child_run)
    return d


def get_first_run_template(para):
    """Return the first <w:r> in a paragraph, used as formatting template."""
    p = para._element
    for child in p:
        if child.tag == qn('w:r'):
            return child
    return None


def clear_runs(para):
    """Remove all <w:r>, <w:ins>, <w:del>, <w:hyperlink> from a paragraph,
       leaving <w:pPr> intact."""
    p = para._element
    for child in list(p):
        if child.tag in (qn('w:pPr'),):
            continue
        p.remove(child)


def replace_with_diff(para, orig_text, new_text):
    """Clear paragraph runs and rebuild as a native diff (<w:ins>/<w:del>/normal runs)."""
    template = get_first_run_template(para)
    a = orig_text.split()
    b = new_text.split()
    sm = difflib.SequenceMatcher(None, a, b)
    clear_runs(para)
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == 'equal':
            txt = ' '.join(b[j1:j2])
            if txt:
                para._element.append(make_run(' ' + txt if para._element.findall(qn('w:r')) else txt,
                                              run_props_template=template))
        elif op == 'insert':
            txt = ' '.join(b[j1:j2])
            if txt:
                prefix = ' ' if has_visible_content(para) else ''
                para._element.append(make_ins(make_run(prefix + txt,
                                                       run_props_template=template)))
        elif op == 'delete':
            txt = ' '.join(a[i1:i2])
            if txt:
                prefix = ' ' if has_visible_content(para) else ''
                para._element.append(make_del(make_del_run(prefix + txt,
                                                           run_props_template=template)))
        elif op == 'replace':
            d_txt = ' '.join(a[i1:i2])
            i_txt = ' '.join(b[j1:j2])
            if d_txt:
                prefix = ' ' if has_visible_content(para) else ''
                para._element.append(make_del(make_del_run(prefix + d_txt,
                                                           run_props_template=template)))
            if i_txt:
                para._element.append(make_ins(make_run(' ' + i_txt,
                                                       run_props_template=template)))


def has_visible_content(para):
    """Check whether paragraph already has any non-pPr child."""
    for child in para._element:
        if child.tag != qn('w:pPr'):
            return True
    return False


def mark_whole_paragraph_as_inserted(para):
    """Wrap the paragraph text in a single <w:ins> and mark the paragraph mark
       as inserted too (this is how Word represents 'whole new paragraph')."""
    template = get_first_run_template(para)
    text = para.text
    clear_runs(para)
    if text:
        para._element.append(make_ins(make_run(text, run_props_template=template)))
    # Mark the paragraph mark itself as inserted: pPr/rPr/ins
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    if rPr.find(qn('w:ins')) is None:
        ins_mark = OxmlElement('w:ins')
        ins_mark.set(qn('w:id'), str(next_id()))
        ins_mark.set(qn('w:author'), AUTHOR)
        ins_mark.set(qn('w:date'), DATE_ISO)
        rPr.append(ins_mark)


def best_match(text, candidates, used, threshold=0.40):
    if len(text) < 20:
        return -1, 0.0
    best_idx, best_ratio = -1, 0.0
    for i, c in enumerate(candidates):
        if i in used or not c:
            continue
        if abs(len(c) - len(text)) > max(len(text), len(c)) * 1.5:
            continue
        r = difflib.SequenceMatcher(None, text.lower(), c.lower()).quick_ratio()
        if r > best_ratio:
            best_idx, best_ratio = i, r
    if best_ratio < threshold or best_idx == -1:
        return -1, best_ratio
    real = difflib.SequenceMatcher(None,
                                   candidates[best_idx].lower(),
                                   text.lower()).ratio()
    return (best_idx, real) if real >= threshold else (-1, real)


def main():
    t0 = time.time()
    orig_doc = Document(ORIG_DOCX)
    orig_paras = [p.text.strip() for p in orig_doc.paragraphs if p.text.strip()]
    for t in orig_doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        orig_paras.append(p.text.strip())
    print(f'Original: {len(orig_paras)} non-empty paragraphs (incl. table cells)')

    rev_doc = Document(REV_DOCX)
    rev_count = sum(1 for p in rev_doc.paragraphs if p.text.strip())
    print(f'Revised:  {rev_count} non-empty body paragraphs (excl. tables)')

    used = set()
    n_eq = n_partial = n_new = 0
    for para in rev_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or '').lower() if para.style else ''
        if 'heading' in style_name or 'title' in style_name:
            continue
        idx, ratio = best_match(text, orig_paras, used)
        if idx >= 0 and ratio >= 0.95:
            used.add(idx)
            n_eq += 1
            continue
        if idx >= 0:
            used.add(idx)
            replace_with_diff(para, orig_paras[idx], text)
            n_partial += 1
        else:
            mark_whole_paragraph_as_inserted(para)
            n_new += 1

    # Append "Deleted from original" block: each orphan original paragraph
    # becomes a new paragraph whose entire content is wrapped in <w:del>.
    deleted = [p for i, p in enumerate(orig_paras) if i not in used]
    if deleted:
        rev_doc.add_page_break()
        title_p = rev_doc.add_paragraph()
        bold_run = title_p.add_run('Deleted content from original manuscript')
        bold_run.bold = True
        for orig_text in deleted:
            new_para = rev_doc.add_paragraph()
            del_run = make_del_run(orig_text)
            new_para._element.append(make_del(del_run))
            # Mark paragraph mark itself as deleted
            pPr = OxmlElement('w:pPr')
            rPr = OxmlElement('w:rPr')
            del_mark = OxmlElement('w:del')
            del_mark.set(qn('w:id'), str(next_id()))
            del_mark.set(qn('w:author'), AUTHOR)
            del_mark.set(qn('w:date'), DATE_ISO)
            rPr.append(del_mark)
            pPr.append(rPr)
            new_para._element.insert(0, pPr)

    if os.path.exists(OUT_DOCX):
        os.remove(OUT_DOCX)
    rev_doc.save(OUT_DOCX)
    size = os.path.getsize(OUT_DOCX)
    print(f'\n  Saved: {size:,} B in {time.time() - t0:.1f}s')
    print(f'  Paragraphs: equal={n_eq}  partial-diff={n_partial}  fully-new={n_new}')
    print(f'  Original paragraphs in deleted-block: {len(deleted)} of {len(orig_paras)}')

    # Verify native track-changes XML
    import zipfile, re
    z = zipfile.ZipFile(OUT_DOCX)
    body = z.read('word/document.xml').decode('utf-8', errors='ignore')
    media = [n for n in z.namelist() if 'media' in n and n.endswith('.png')]
    ins_count = len(re.findall(r'<w:ins[\s>]', body))
    del_count = len(re.findall(r'<w:del[\s>]', body))
    print(f'  PNGs preserved:           {len(media)}')
    print(f'  <w:ins> elements (TRUE):  {ins_count:,}')
    print(f'  <w:del> elements (TRUE):  {del_count:,}')


if __name__ == '__main__':
    main()
