from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
MD_PATH = REPO_ROOT / "docs" / "proposals" / "EPC_MAINTENANCE_AI_PROJECT_BRIEF_V0.1.ko.md"
ASSET_DIR = REPO_ROOT / "docs" / "proposals" / "assets" / "project-intro"
DOCX_PATH = REPO_ROOT / "docs" / "proposals" / "EPC_MAINTENANCE_AI_PROJECT_BRIEF_V0.1.ko.docx"


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    if bold:
        candidates = [
            "C:/Windows/Fonts/segoeuib.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
            *candidates,
        ]
    for font_path in candidates:
        fp = Path(font_path)
        if fp.exists():
            return ImageFont.truetype(str(fp), size=size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=4)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def draw_box(
    draw: ImageDraw.ImageDraw,
    rect: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: tuple[int, int, int],
    border: tuple[int, int, int] = (45, 55, 72),
) -> None:
    draw.rounded_rectangle(rect, radius=16, fill=fill, outline=border, width=3)
    x1, y1, x2, y2 = rect
    title_font = load_font(32, bold=True)
    body_font = load_font(24)

    draw.text((x1 + 18, y1 + 14), title, fill=(15, 26, 42), font=title_font)
    draw.line((x1 + 16, y1 + 58, x2 - 16, y1 + 58), fill=(123, 141, 163), width=2)

    max_width = (x2 - x1) - 36
    wrapped = wrap_text(draw, body, body_font, max_width)
    draw.multiline_text((x1 + 18, y1 + 72), wrapped, fill=(22, 34, 51), font=body_font, spacing=6)


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: tuple[int, int, int] = (55, 95, 165),
    width: int = 6,
) -> None:
    draw.line((start, end), fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        tip = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        tip = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(tip, fill=color)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> str:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        trial = " ".join([*current, word])
        width, _ = text_size(draw, trial, font)
        if width <= max_width:
            current.append(word)
        else:
            if current:
                lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return "\n".join(lines)


def create_system_architecture(path: Path) -> None:
    image = Image.new("RGB", (1800, 1020), (248, 251, 255))
    draw = ImageDraw.Draw(image)

    header_font = load_font(42, bold=True)
    draw.text((42, 24), "System Architecture Overview", fill=(21, 34, 56), font=header_font)

    boxes = [
        ((50, 120, 370, 360), "Input Layer", "Project + Asset + Discipline fields\n3-pane workbench UI", (225, 239, 255)),
        ((430, 120, 770, 360), "Calculation API", "Typed request/response contract\nMock / Backend mode switching", (222, 245, 230)),
        ((830, 120, 1170, 360), "Domain Engines", "7 discipline services\nPiping/Vessel/Rotating/.../Civil", (255, 241, 220)),
        ((1230, 120, 1570, 360), "Verification", "4-layer gates\nfail-closed red-flag policy", (255, 227, 229)),
    ]

    for rect, title, body, fill in boxes:
        draw_box(draw, rect, title, body, fill)

    draw_arrow(draw, (370, 240), (430, 240))
    draw_arrow(draw, (770, 240), (830, 240))
    draw_arrow(draw, (1170, 240), (1230, 240))

    draw_box(
        draw,
        (250, 430, 930, 730),
        "Regulatory & Knowledge Layer",
        "KOSHA API + local snapshots\nlaw_articles + guide_documents + local RAG",
        (233, 240, 255),
    )
    draw_box(
        draw,
        (1020, 430, 1750, 730),
        "Output & Action Layer",
        "Summary metrics + standards trace + legal mapping\nblocked banner + recommendations + export",
        (230, 248, 242),
    )

    draw_arrow(draw, (590, 430), (590, 360))
    draw_arrow(draw, (1360, 430), (1360, 360))
    draw_arrow(draw, (930, 580), (1020, 580))

    note_font = load_font(22)
    draw.text(
        (42, 920),
        "Design principle: traceability + verification + safety-blocking in one workflow.",
        fill=(64, 80, 104),
        font=note_font,
    )

    image.save(path)


def create_user_journey(path: Path) -> None:
    image = Image.new("RGB", (1800, 980), (250, 252, 255))
    draw = ImageDraw.Draw(image)
    draw.text((42, 24), "User Journey (Operator Workflow)", fill=(23, 38, 62), font=load_font(42, bold=True))

    steps = [
        ("1. Enter", "Open dashboard\nselect discipline"),
        ("2. Input", "Fill form or preset\nvalidate fields"),
        ("3. Run", "Execute calculation\nget key metrics"),
        ("4. Verify", "Check 4 layers\nreview flags"),
        ("5. Evidence", "Review standards\nand legal mapping"),
        ("6. Share", "Export JSON/MD\nreport to team"),
    ]

    x = 70
    y = 220
    box_w = 250
    box_h = 190
    gap = 35

    for idx, (title, body) in enumerate(steps):
        rect = (x + idx * (box_w + gap), y, x + idx * (box_w + gap) + box_w, y + box_h)
        draw_box(draw, rect, title, body, (233, 244, 255))
        if idx < len(steps) - 1:
            start = (rect[2], y + box_h // 2)
            end = (rect[2] + gap, y + box_h // 2)
            draw_arrow(draw, start, end, color=(71, 111, 179), width=5)

    draw_box(
        draw,
        (120, 520, 1680, 830),
        "Key Outcome",
        "Fast and auditable decision-making: calculation + verification + standards/legal evidence are linked in one place.",
        (225, 248, 232),
    )
    image.save(path)


def create_roadmap(path: Path) -> None:
    image = Image.new("RGB", (1800, 980), (252, 252, 252))
    draw = ImageDraw.Draw(image)
    draw.text((42, 24), "Growth Roadmap (6M / 1Y / 2Y)", fill=(25, 33, 47), font=load_font(42, bold=True))

    columns = [
        ("6 Months", "Pilot stability\ncore disciplines deepening\ninitial paid pilots"),
        ("1 Year", "Regulatory automation\nmulti-site rollout\nAPI enterprise integration"),
        ("2 Years", "Policy engine maturity\npartner channels\nindustry package expansion"),
    ]
    starts = [130, 665, 1200]
    colors = [(226, 239, 255), (228, 247, 233), (255, 239, 221)]

    for idx, (title, body) in enumerate(columns):
        x1 = starts[idx]
        rect = (x1, 210, x1 + 470, 760)
        draw_box(draw, rect, title, body, colors[idx])
        if idx < 2:
            draw_arrow(draw, (rect[2], 485), (rect[2] + 45, 485), color=(83, 108, 156), width=7)

    draw.line((95, 820, 1710, 820), fill=(128, 143, 166), width=4)
    for marker_x in [230, 760, 1300]:
        draw.ellipse((marker_x - 11, 809, marker_x + 11, 831), fill=(65, 93, 150), outline=(65, 93, 150))

    draw.text((205, 845), "Q1-Q2", fill=(50, 66, 94), font=load_font(24, bold=True))
    draw.text((735, 845), "Q3-Q4", fill=(50, 66, 94), font=load_font(24, bold=True))
    draw.text((1275, 845), "Year 2+", fill=(50, 66, 94), font=load_font(24, bold=True))

    image.save(path)


def generate_diagrams() -> None:
    ensure_dir(ASSET_DIR)
    create_system_architecture(ASSET_DIR / "system_architecture.png")
    create_user_journey(ASSET_DIR / "user_journey.png")
    create_roadmap(ASSET_DIR / "roadmap_timeline.png")


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table_rows(lines: Iterable[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def is_separator_row(row: list[str]) -> bool:
    if not row:
        return False
    for cell in row:
        token = cell.replace(":", "").replace("-", "").strip()
        if token:
            return False
    return True


def add_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            title = heading.group(2).strip()
            if title.startswith("PART 2."):
                doc.add_page_break()
            doc.add_heading(title, level=level)
            i += 1
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            alt = image_match.group(1).strip()
            rel_path = image_match.group(2).strip()
            img_path = (md_path.parent / rel_path).resolve()
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.5))
                if alt:
                    caption = doc.add_paragraph(alt)
                    caption.style = "Caption"
            i += 1
            continue

        if is_table_line(stripped):
            table_block: list[str] = []
            while i < len(lines) and is_table_line(lines[i].strip()):
                table_block.append(lines[i])
                i += 1
            rows = parse_table_rows(table_block)
            if not rows:
                continue

            header = rows[0]
            data_rows = rows[1:]
            if data_rows and is_separator_row(data_rows[0]):
                data_rows = data_rows[1:]

            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for col_idx, text in enumerate(header):
                table.rows[0].cells[col_idx].text = text
            for row in data_rows:
                values = row + [""] * (len(header) - len(row))
                cells = table.add_row().cells
                for col_idx in range(len(header)):
                    cells[col_idx].text = values[col_idx]
            continue

        if re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(content, style="List Number")
            i += 1
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        doc.add_paragraph(stripped)
        i += 1

    ensure_dir(docx_path.parent)
    doc.save(docx_path)


def main() -> None:
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Markdown source not found: {MD_PATH}")
    generate_diagrams()
    add_markdown_to_docx(MD_PATH, DOCX_PATH)
    print(f"Generated diagrams in: {ASSET_DIR}")
    print(f"Generated docx: {DOCX_PATH}")


if __name__ == "__main__":
    main()
